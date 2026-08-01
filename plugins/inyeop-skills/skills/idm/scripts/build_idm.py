# -*- coding: utf-8 -*-
"""FV IDM(Initial Deal Memo) 생성기. data JSON -> DOCX.

사용법:
    python3 build_idm.py data.json "회사명 IDM.docx"

data JSON 스키마 (example_data.json 참고):
{
  "company": "회사명 (영문)",            # 문서 제목 -> "회사명 IDM"
  "table": [                           # 7행 x [좌라벨, 좌값, 우라벨, 우값]
    ["분야", "...", "투자유형", "..."],
    ...
  ],
  "sections": [
    {"title": "Deal Structure", "items": [ {"t":"...", "lv":0}, ... ]},
    {"title": "Good Points", "ordered": true,
        "items": [ {"t":"...", "lv":0, "bold": true}, {"t":"...", "lv":1} ]},
    {"title": "Risk Factors", "ordered": true, "accent": true, "items": [...]}
  ]
}

규칙:
- lv 0/1/2 -> 불릿 글머리 • / ◦ / ▪. ordered:true 섹션의 lv0 은 "1)" 자동 증가.
- 모르는 표 값은 빈 문자열("")로 두면 공란으로 렌더(추측 금지 원칙).
- 아이템 텍스트에 " : " 가 있으면 앞부분(라벨)을 자동 굵게.
- accent:true 섹션은 ordered lv0 라벨을 코랄레드로 강조(Risk Factors 용).
"""
import json, sys
from docx import Document
from docx.shared import Pt, RGBColor, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "맑은 고딕"          # 사용자 Windows Word 기준 한국어 폰트
INK = RGBColor(0x1A, 0x1A, 0x2E)
GREY = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
LABEL_SHADE = "EEF1F5"
BORDER = "CCCCCC"
RULE = "BBBBBB"
COLW = [1500, 3180, 1680, 3000]      # twips, 합 9360 (US Letter, 1" 여백)


def _el(t):
    return OxmlElement(t)


def cjk(run, name=FONT):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = _el('w:rFonts'); rPr.insert(0, rFonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(a), name)


def add_run(p, text, size=10, bold=False, color=INK):
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = color
    cjk(r)
    return r


def tight(p, before=2, after=4, line=1.06):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line


def bottom_border(p, color=RULE, sz=6):
    pPr = p._p.get_or_add_pPr()
    pbdr = _el('w:pBdr'); bottom = _el('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), color)
    pbdr.append(bottom); pPr.append(pbdr)


# ---------- table helpers ----------
def set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = _el('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_table_borders(table, sz=4, color=BORDER):
    b = _el('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = _el('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color); b.append(e)
    table._tbl.tblPr.append(b)


def set_fixed_layout(table, total_tw):
    tblPr = table._tbl.tblPr
    lay = _el('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)
    w = _el('w:tblW'); w.set(qn('w:w'), str(total_tw)); w.set(qn('w:type'), 'dxa'); tblPr.append(w)


def set_grid(table, widths_tw):
    tbl = table._tbl; grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = _el('w:tblGrid'); tbl.insert(0, grid)
    for gc in grid.findall(qn('w:gridCol')):
        grid.remove(gc)
    for w in widths_tw:
        gc = _el('w:gridCol'); gc.set(qn('w:w'), str(w)); grid.append(gc)


def set_cell_width(cell, tw):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    w = _el('w:tcW'); w.set(qn('w:w'), str(tw)); w.set(qn('w:type'), 'dxa'); tcPr.append(w)


def fill_cell(cell, text, label=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tight(p, 1, 1, 1.05)
    add_run(p, text or "", size=9.5, bold=label, color=INK)
    if label:
        set_cell_bg(cell, LABEL_SHADE)


# ---------- bullet items ----------
PREFIX = {0: "•", 1: "◦", 2: "▪"}
INDENT = {0: 360, 1: 720, 2: 1080}   # twips left
HANG = 250


def add_item(doc, it, ordered=False, accent=False, counter=None):
    lv = int(it.get("lv", 0))
    text = it.get("t", "")
    whole_bold = bool(it.get("bold", False))
    left = INDENT.get(lv, 360)
    p = doc.add_paragraph()
    tight(p, 2, 4, 1.07)
    pf = p.paragraph_format
    pf.left_indent = Twips(left); pf.first_line_indent = Twips(-HANG)
    p.paragraph_format.tab_stops.add_tab_stop(Twips(left), WD_TAB_ALIGNMENT.LEFT)

    if ordered and lv == 0 and counter is not None:
        pref = "%d)" % counter
        pcolor = ACCENT if accent else INK
        add_run(p, pref + "\t", size=10, bold=True, color=pcolor)
    else:
        add_run(p, PREFIX.get(lv, "•") + "\t", size=10, bold=False, color=INK)

    color = ACCENT if (accent and ordered and lv == 0) else INK
    if whole_bold:
        add_run(p, text, size=10, bold=True, color=color)
    elif " : " in text:
        i = text.index(" : ")
        add_run(p, text[:i], size=10, bold=True, color=INK)
        add_run(p, text[i:], size=10, bold=False, color=INK)
    else:
        add_run(p, text, size=10, bold=False, color=INK)


def main():
    data_path, out_path = sys.argv[1], sys.argv[2]
    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(10)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = _el('w:rFonts'); rpr.insert(0, rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(a), FONT)

    pt = doc.add_paragraph(); tight(pt, 0, 2, 1.0)
    add_run(pt, "%s IDM" % data.get("company", ""), size=19, bold=True, color=INK)
    ps = doc.add_paragraph(); tight(ps, 0, 10, 1.0)
    add_run(ps, "Initial Deal Memo", size=11, bold=False, color=GREY)

    ph = doc.add_paragraph(); tight(ph, 6, 6, 1.0); bottom_border(ph)
    add_run(ph, "1. Company Name", size=13, bold=True, color=INK)

    rows = data.get("table", [])
    table = doc.add_table(rows=len(rows), cols=4)
    set_fixed_layout(table, sum(COLW)); set_grid(table, COLW); set_table_borders(table)
    for ri, row in enumerate(rows):
        cells = table.rows[ri].cells
        vals = (list(row) + ["", "", "", ""])[:4]
        for ci in range(4):
            set_cell_width(cells[ci], COLW[ci])
            fill_cell(cells[ci], vals[ci], label=(ci % 2 == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    for s in data.get("sections", []):
        ph = doc.add_paragraph(); tight(ph, 12, 6, 1.0); bottom_border(ph)
        add_run(ph, s.get("title", ""), size=13, bold=True, color=INK)
        ordered = bool(s.get("ordered", False))
        accent = bool(s.get("accent", False))
        counter = 0
        for it in s.get("items", []):
            if ordered and int(it.get("lv", 0)) == 0:
                counter += 1
                add_item(doc, it, ordered=True, accent=accent, counter=counter)
            else:
                add_item(doc, it, ordered=ordered, accent=accent, counter=None)

    doc.save(out_path)
    print("saved:", out_path)


if __name__ == "__main__":
    main()
