# -*- coding: utf-8 -*-
"""스타트업플러스 온라인 투자밋업 보고서 생성기. data dict -> DOCX -> PDF."""
import os, subprocess, sys
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT="Noto Sans CJK KR"; GRAY_SECTION="BFBFBF"; GRAY_FIELD="F2F2F2"
FIXED={"org":"패스트벤처스 주식회사","manager":"박인엽 팀장",
 "biz_field":"IT/운영관리/라이프스타일/미디어/엔터테인먼트/바이오/에너지/프롭테크/하드웨어",
 "tech_field":"AI/데이터/로보틱스/사물인터넷/신소재/전고체배터리/제조/클라우드",
 "target":"무관","funds":["패스트 Core-1 투자조합","패스트 2022 Seed 투자조합"],
 "meet_type":"온라인 미팅","account":"우리은행 / 1002-288-930613"}
COLW_PT=[61.0,69.7,42.5,90.3,79.4,192.1]
def pt2tw(p): return int(round(p*20))
COLW_TW=[pt2tw(w) for w in COLW_PT]; TBL_TW=sum(COLW_TW)
def _el(t): return OxmlElement(t)
def set_cell_bg(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=_el('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def set_table_borders(table,sz=4,color='000000'):
    b=_el('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=_el('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),str(sz)); e.set(qn('w:space'),'0'); e.set(qn('w:color'),color); b.append(e)
    table._tbl.tblPr.append(b)
def set_fixed_layout(table,total_tw):
    tblPr=table._tbl.tblPr
    lay=_el('w:tblLayout'); lay.set(qn('w:type'),'fixed'); tblPr.append(lay)
    w=_el('w:tblW'); w.set(qn('w:w'),str(total_tw)); w.set(qn('w:type'),'dxa'); tblPr.append(w)
def set_cell_margins(table,top=18,bottom=18,left=72,right=72):
    mar=_el('w:tblCellMar')
    for k,v in (('top',top),('left',left),('bottom',bottom),('right',right)):
        e=_el('w:'+k); e.set(qn('w:w'),str(v)); e.set(qn('w:type'),'dxa'); mar.append(e)
    table._tbl.tblPr.append(mar)
def set_one_cell_margins(cell,top=0,bottom=0,left=72,right=72):
    tcPr=cell._tc.get_or_add_tcPr(); mar=_el('w:tcMar')
    for k,v in (('top',top),('left',left),('bottom',bottom),('right',right)):
        e=_el('w:'+k); e.set(qn('w:w'),str(v)); e.set(qn('w:type'),'dxa'); mar.append(e)
    tcPr.append(mar)
def set_grid(table,widths_tw):
    tbl=table._tbl; grid=tbl.find(qn('w:tblGrid'))
    if grid is None: grid=_el('w:tblGrid'); tbl.insert(0,grid)
    for gc in grid.findall(qn('w:gridCol')): grid.remove(gc)
    for w in widths_tw:
        gc=_el('w:gridCol'); gc.set(qn('w:w'),str(w)); grid.append(gc)
def set_cell_width(cell,tw):
    tcPr=cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    w=_el('w:tcW'); w.set(qn('w:w'),str(tw)); w.set(qn('w:type'),'dxa'); tcPr.append(w)
def cjk(run,name=FONT):
    run.font.name=name; rPr=run._element.get_or_add_rPr(); rFonts=rPr.find(qn('w:rFonts'))
    if rFonts is None: rFonts=_el('w:rFonts'); rPr.insert(0,rFonts)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rFonts.set(qn(a),name)
ALIGN={'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER}
def _tight(p):
    pf=p.paragraph_format; pf.space_before=Pt(0); pf.space_after=Pt(0); pf.line_spacing=1.0
def fill_cell(cell,lines,size=10,bold=False,align='left',valign='center',bg=None):
    if isinstance(lines,str): lines=[lines]
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER if valign=='center' else WD_CELL_VERTICAL_ALIGNMENT.TOP
    if bg: set_cell_bg(cell,bg)
    for extra in cell.paragraphs[1:]: extra._element.getparent().remove(extra._element)
    p0=cell.paragraphs[0]
    for r in list(p0.runs): r._element.getparent().remove(r._element)
    for i,txt in enumerate(lines):
        p=p0 if i==0 else cell.add_paragraph(); p.alignment=ALIGN[align]; _tight(p)
        run=p.add_run(txt); run.font.size=Pt(size); run.font.bold=bold; cjk(run)
    return cell
def add_center_tabs(p,centers):
    pPr=p._p.get_or_add_pPr(); tabs=_el('w:tabs')
    for c in centers:
        t=_el('w:tab'); t.set(qn('w:val'),'center'); t.set(qn('w:pos'),str(int(c))); tabs.append(t)
    pPr.append(tabs)
def checkbox_row(cell,items,span_tw,checked_idx,two_line=False,descs=None,size_main=9,size_desc=8.5):
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    n=len(items); seg=span_tw/n; centers=[seg*(i+0.5) for i in range(n)]
    for extra in cell.paragraphs[1:]: extra._element.getparent().remove(extra._element)
    p1=cell.paragraphs[0]
    for r in list(p1.runs): r._element.getparent().remove(r._element)
    p1.alignment=ALIGN['left']; _tight(p1); add_center_tabs(p1,centers)
    box=lambda i:"■" if i==checked_idx else "□"
    r=p1.add_run("".join("\t"+box(i)+" "+items[i] for i in range(n))); r.font.size=Pt(size_main); cjk(r)
    if two_line and descs:
        p2=cell.add_paragraph(); p2.alignment=ALIGN['left']; _tight(p2); add_center_tabs(p2,centers)
        r2=p2.add_run("".join("\t"+descs[i] for i in range(n))); r2.font.size=Pt(size_desc); cjk(r2)
def strip_cell(cell):
    for p in list(cell.paragraphs): p._element.getparent().remove(p._element)
def add_trailing_p(cell):
    tc=cell._tc; p=_el('w:p'); pPr=_el('w:pPr')
    sp=_el('w:spacing'); sp.set(qn('w:line'),'16'); sp.set(qn('w:lineRule'),'exact'); sp.set(qn('w:after'),'0'); pPr.append(sp)
    rPr=_el('w:rPr'); sz=_el('w:sz'); sz.set(qn('w:val'),'2'); rPr.append(sz); pPr.append(rPr)
    p.append(pPr); tc.append(p)
def minimize_para(p_el):
    pPr=p_el.find(qn('w:pPr'))
    if pPr is None: pPr=_el('w:pPr'); p_el.insert(0,pPr)
    sp=_el('w:spacing'); sp.set(qn('w:before'),'0'); sp.set(qn('w:after'),'0'); sp.set(qn('w:line'),'16'); sp.set(qn('w:lineRule'),'exact'); pPr.append(sp)
    rPr=_el('w:rPr'); sz=_el('w:sz'); sz.set(qn('w:val'),'2'); rPr.append(sz); pPr.append(rPr)
def photo_block(cell,span_tw,photos,caption="밋업 진행 사진"):
    half=span_tw//2; set_one_cell_margins(cell,0,0,0,0)
    nt=cell.add_table(rows=2,cols=2)
    set_grid(nt,[half,span_tw-half]); set_fixed_layout(nt,span_tw)
    set_table_borders(nt,sz=4); set_cell_margins(nt,top=2,bottom=2,left=6,right=6); nt.alignment=WD_TABLE_ALIGNMENT.CENTER
    from PIL import Image as _PIL
    max_w=(half/20)-16; max_h=108
    for c in range(2):
        ic=nt.cell(0,c); set_cell_width(ic, half if c==0 else span_tw-half); ic.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=ic.paragraphs[0]; p.alignment=ALIGN['center']; _tight(p)
        if photos and c<len(photos) and photos[c] and os.path.exists(photos[c]):
            _iw,_ih=_PIL.open(photos[c]).size; _sc=min(max_w/_iw, max_h/_ih)
            p.add_run().add_picture(photos[c],width=Pt(_iw*_sc),height=Pt(_ih*_sc))
        else:
            rr=p.add_run("(사진)"); rr.font.size=Pt(9); cjk(rr)
        cc=nt.cell(1,c); set_cell_width(cc, half if c==0 else span_tw-half); fill_cell(cc,caption,size=10,align='center')
    tc=cell._tc
    for pe in tc.findall(qn('w:p')):
        if not pe.findall(qn('w:r')): tc.remove(pe)
    add_trailing_p(cell)
def vlabel(cell,lines,size=9):
    fill_cell(cell,lines,size=size,bold=True,align='center')
    set_one_cell_margins(cell,top=6,bottom=6,left=40,right=40)
def build(data,out_docx):
    doc=Document(); st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(10)
    st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
    sec=doc.sections[0]; sec.page_width=Mm(210); sec.page_height=Mm(297)
    sec.left_margin=Pt(30); sec.right_margin=Pt(30); sec.top_margin=Pt(16); sec.bottom_margin=Pt(16)
    title=doc.add_paragraph(); title.alignment=ALIGN['center']; _tight(title); title.paragraph_format.space_after=Pt(2)
    tr=title.add_run("<스타트업플러스 온라인 투자밋업 보고서>"); tr.font.size=Pt(17); tr.bold=True; cjk(tr)
    tbl=doc.add_table(rows=16,cols=6); set_grid(tbl,COLW_TW); set_fixed_layout(tbl,TBL_TW)
    set_table_borders(tbl,sz=4); set_cell_margins(tbl,top=6,bottom=6,left=72,right=72); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for r in range(16):
        for c in range(6): set_cell_width(tbl.cell(r,c),COLW_TW[c])
    def C(r,c): return tbl.cell(r,c)
    def hm(r,c0,c1):
        a=C(r,c0)
        for c in range(c0+1,c1+1): a=a.merge(C(r,c))
        return a
    fill_cell(C(0,0),"밋업정보",align='center',bg=GRAY_SECTION)
    fill_cell(C(0,1),"밋업일시",align='center',bg=GRAY_FIELD)
    fill_cell(hm(0,2,3),data["meet_date"],align='center')
    fill_cell(C(0,4),"밋업형식",align='center',bg=GRAY_FIELD)
    fill_cell(C(0,5),FIXED["meet_type"],align='center')
    fill_cell(C(1,0),["투자자","정보"],align='center',bg=GRAY_SECTION)
    fill_cell(C(1,1),"투자기관명",align='center',bg=GRAY_FIELD)
    fill_cell(hm(1,2,3),FIXED["org"],align='center')
    fill_cell(C(1,4),["담당자 성함","및 직함"],align='center',bg=GRAY_FIELD)
    fill_cell(C(1,5),FIXED["manager"],align='center')
    fill_cell(C(2,1),"관심사업분야",align='center',bg=GRAY_FIELD)
    fill_cell(hm(2,2,3),FIXED["biz_field"],size=9,align='left')
    fill_cell(C(2,4),"관심기술분야",align='center',bg=GRAY_FIELD)
    fill_cell(C(2,5),FIXED["tech_field"],size=9,align='left')
    fill_cell(C(3,1),["투자밋업","희망대상"],align='center',bg=GRAY_FIELD)
    fill_cell(hm(3,2,3),FIXED["target"],align='center')
    fill_cell(C(3,4),"보유펀드명",align='center',bg=GRAY_FIELD)
    fill_cell(C(3,5),FIXED["funds"],align='center')
    fill_cell(C(4,1),"투자기관구분",align='center',bg=GRAY_FIELD)
    checkbox_row(hm(4,2,5),["창업기획자","기술지주회사","창업투자회사","유한책임회사","기타"],sum(COLW_TW[2:6]),checked_idx=2,size_main=10)
    C(1,0).merge(C(4,0))
    fill_cell(C(5,0),["스타트업","정보"],align='center',bg=GRAY_SECTION)
    fill_cell(C(5,1),"기업명",align='center',bg=GRAY_FIELD)
    fill_cell(hm(5,2,3),data["company"],align='center')
    fill_cell(C(5,4),"담당자명/직급",align='center',bg=GRAY_FIELD)
    fill_cell(C(5,5),data["ceo"],align='center')
    fill_cell(C(6,0),["밋업","내용"],align='center',bg=GRAY_SECTION)
    fill_cell(C(6,1),["사업","아이템"],align='center',bg=GRAY_FIELD)
    vlabel(C(6,2),["아이템","개요"])
    fill_cell(hm(6,3,5),data["item_overview"],align='left')
    vlabel(C(7,2),["차별화","특징"])
    fill_cell(hm(7,3,5),data["diff_feature"],align='left')
    C(6,1).merge(C(7,1))
    fill_cell(C(8,1),["투자검토","의견"],align='center',bg=GRAY_FIELD)
    vlabel(C(8,2),["시장","적합성"])
    fill_cell(hm(8,3,5),data["market_fit"],align='left')
    vlabel(C(9,2),["기술","차별성"])
    fill_cell(hm(9,3,5),data["tech_diff"],align='left')
    vlabel(C(10,2),["재무","현황"])
    fill_cell(hm(10,3,5),data["finance"],align='left')
    C(8,1).merge(C(10,1))
    fill_cell(C(11,1),["종합검토","의견"],align='center',bg=GRAY_FIELD)
    vlabel(C(11,2),["투자","의향"])
    checkbox_row(hm(11,3,5),["5점","4점","3점","2점","1점"],sum(COLW_TW[3:6]),checked_idx=(5-int(data["intent"])),two_line=True,
                 descs=["(투자의향 매우높음)","(투자의향 높음)","(투자의향 보통)","(투자의향 낮음)","(투자의향 매우낮음)"])
    vlabel(C(12,2),["긍정적","검토의견"])
    fill_cell(hm(12,3,5),data["positive"],align='left')
    vlabel(C(13,2),["부정적","검토의견"])
    fill_cell(hm(13,3,5),data["negative"],align='left')
    C(11,1).merge(C(13,1)); C(6,0).merge(C(13,0))
    fill_cell(C(14,0),["밋업","사진"],align='center',bg=GRAY_SECTION)
    pc=hm(14,1,5); strip_cell(pc); photo_block(pc,sum(COLW_TW[1:6]),data.get("photos"))
    fill_cell(C(15,0),["계좌 번호","(은행명/","계좌)"],align='center',bg=GRAY_SECTION)
    fill_cell(hm(15,1,5),FIXED["account"],align='center')
    tp=doc.add_paragraph(); minimize_para(tp._p)
    doc.save(out_docx); return out_docx
def to_pdf(docx_path,outdir):
    subprocess.run(["soffice","--headless","--convert-to","pdf","--outdir",outdir,docx_path],check=True,capture_output=True,timeout=120)
    return os.path.join(outdir,os.path.splitext(os.path.basename(docx_path))[0]+".pdf")
if __name__=="__main__":
    import json
    d=json.load(open(sys.argv[1],encoding="utf-8")); o=sys.argv[2]
    build(d,o); print("DOCX:",o); print("PDF:",to_pdf(o,os.path.dirname(o) or "."))
