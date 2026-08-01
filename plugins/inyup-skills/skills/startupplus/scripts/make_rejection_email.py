# -*- coding: utf-8 -*-
"""
투자 거절 메일 생성기 (박인엽/패스트벤처스 템플릿).
입력 JSON: 단일 dict 또는 dict 리스트. 각 dict 키:
  company   : 회사명 (메일 본문 'OOOOO 사'에 사용)
  ceo       : 대표명 (직함 '대표'는 자동 부가)
  story     : 도메인 스토리 ('대표님의 OOO에 대한 재밌는 스토리'의 OOO)
  strength  : 강점 1개 ('(b) ...'에 들어갈 문구. 보통 시장 매력/핵심 강점)
  negative  : 거절 사유 문장 (단점 2~3개를 녹인 한 문장, '다만, ...관점에서 내부 설득이 잘 이루어지지 못 했던 것 같습니다.')
사용: python3 make_rejection_email.py email_data.json /출력/폴더
출력: '투자 거절 메일(회사명).docx' (회사별 1개)
"""
import sys, os, json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
FONT="Noto Sans CJK KR"
def cjk(run,b=False,size=11):
    run.font.name=FONT; run.font.size=Pt(size); run.font.bold=b
    rPr=run._element.get_or_add_rPr(); rf=rPr.find(qn('w:rFonts'))
    if rf is None: rf=run._element.makeelement(qn('w:rFonts'),{}); rPr.insert(0,rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rf.set(qn(a),FONT)
def make_one(d, outdir):
    co=d["company"]; ceo=d["ceo"]; story=d["story"]; strength=d["strength"]; neg=d["negative"]
    doc=Document(); st=doc.styles['Normal']; st.font.name=FONT; st.font.size=Pt(11)
    st.element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
    def line(txt,after=2):
        p=doc.add_paragraph(); pf=p.paragraph_format
        pf.space_before=Pt(0); pf.space_after=Pt(after); pf.line_spacing=1.3
        r=p.add_run(txt); cjk(r); return p
    line(f"안녕하세요 {ceo} 대표님, 패스트벤처스의 박인엽입니다.")
    line("귀한 시간 내주시고 회사와 사업에 대해 소개 주셔서 다시 한 번 감사 드립니다.")
    line(f"대표님의 {story}에 대한 재밌는 스토리에 많이 공감이 되었습니다.")
    line("저희쪽에서 빠르게 투자 검토에 대한 의견을 전달 드리는 것을 선호하실 것 같아, 충분한 시간을 두고 오래오래 심사숙고하기 보다는, 최대한 빠르게 내부 논의를 통하여 결론을 내어보았습니다.")
    line(f"결과적으로, 저희 패스트벤처스는 이번에 {co} 사에 대한 투자를 진행하기 조금 어려울 것 같습니다.")
    line(f"(a) 대표님의 팀을 운영하시는 전문성, (b) {strength} 등 긍정적인 요소들이 많다고 느껴졌기 때문에 더 고민이 되었는데요,")
    line(neg)
    line("본 검토 의견은, 저희 패스트벤처스가 가지는 관점의 차이에서 기인한 것일 뿐이고, 회사에 대해 가지는 확신의 크기가 아주 조금 부족했기 때문이지, 절대 회사와 사업 자체에 대해 부정적으로 생각하여 투자를 못 하게 된 것은 아님을 꼭 말씀 드리고 싶습니다.")
    line("다음에 다른 기회로 또 인사 드릴 수 있으면 좋겠습니다.")
    line("감사합니다.")
    line("박인엽 드림.", after=0)
    out=os.path.join(outdir, f"투자 거절 메일({co}).docx")
    doc.save(out); return out
if __name__=="__main__":
    data=json.load(open(sys.argv[1],encoding="utf-8"))
    outdir=sys.argv[2] if len(sys.argv)>2 else "."
    items=data if isinstance(data,list) else [data]
    for d in items:
        print("SAVED:", make_one(d,outdir))
